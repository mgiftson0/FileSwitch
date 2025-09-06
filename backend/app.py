from flask import Flask, request, jsonify, send_file
import os
from werkzeug.utils import secure_filename
from pdf2docx import Converter as PdfToDocxConverter
from flask_cors import CORS
import logging
import re
from bs4 import BeautifulSoup

app = Flask(__name__)

# CORS configuration
CORS(app, resources={
    r"/*": {
        "origins": ["http://localhost:5173", "http://127.0.0.1:5173", "https://file-switch-5u3h.vercel.app/"],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"]
    }
})

UPLOAD_FOLDER = 'uploads'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file part'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400
        
        conversion_type = request.form.get('type', 'pdf_to_word')
        print(f"Conversion type: {conversion_type}")
        
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        print(f"File saved: {file_path}")

        # Handle different conversion types
        if conversion_type == 'pdf_to_word':
            converted_file_path = convert_pdf_to_word(file_path)
        elif conversion_type == 'word_to_pdf':
            converted_file_path = convert_word_to_pdf(file_path)
        elif conversion_type == 'html_to_pdf':
            converted_file_path = convert_html_to_pdf_perfect(file_path)
        elif conversion_type == 'text_to_docx':
            converted_file_path = convert_html_to_docx_perfect(file_path)
        else:
            return jsonify({'error': 'Invalid conversion type'}), 400

        converted_filename = os.path.basename(converted_file_path)
        print(f"Converted file: {converted_filename}")
        
        return jsonify({
            'original_id': filename,
            'converted_id': converted_filename
        }), 200
        
    except Exception as e:
        print(f"Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Conversion failed: {str(e)}'}), 500

def convert_pdf_to_word(file_path):
    docx_file = file_path.replace('.pdf', '.docx')
    cv = PdfToDocxConverter(file_path)
    cv.convert(docx_file)
    cv.close()
    return docx_file

def convert_word_to_pdf(file_path):
    """Convert Word document to PDF using python-docx and reportlab"""
    try:
        from docx import Document
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.lib import colors
        
        print(f"Converting Word to PDF: {file_path}")
        
        # Load Word document
        doc = Document(file_path)
        
        # Create PDF file path
        pdf_file = file_path.replace('.docx', '.pdf').replace('.doc', '.pdf')
        
        # Create PDF document
        pdf_doc = SimpleDocTemplate(pdf_file, pagesize=letter, 
                                  rightMargin=72, leftMargin=72, 
                                  topMargin=72, bottomMargin=72)
        
        # Create styles
        styles = getSampleStyleSheet()
        create_perfect_styles(styles)
        
        story = []
        
        # Process Word document paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Determine style based on paragraph style
                style_name = 'CustomNormal'
                
                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    try:
                        level = paragraph.style.name.replace('Heading ', '').strip()
                        if level.isdigit() and 1 <= int(level) <= 6:
                            style_name = f'CustomH{level}'
                    except:
                        style_name = 'CustomH1'
                
                # Convert paragraph with formatting
                formatted_text = convert_docx_paragraph_to_reportlab(paragraph)
                
                if formatted_text.strip():
                    story.append(Paragraph(formatted_text, styles[style_name]))
                    story.append(Spacer(1, 6))
        
        # Process tables if any
        for table in doc.tables:
            story.append(Spacer(1, 12))
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    story.append(Paragraph(row_text, styles['CustomNormal']))
            story.append(Spacer(1, 12))
        
        if not story:
            story.append(Paragraph("Document appears to be empty", styles['CustomNormal']))
        
        print(f"Building PDF with {len(story)} elements...")
        pdf_doc.build(story)
        print(f"Word to PDF conversion successful: {pdf_file}")
        
        return pdf_file
        
    except Exception as e:
        print(f"Word to PDF conversion error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Word to PDF conversion failed: {str(e)}")

def convert_docx_paragraph_to_reportlab(paragraph):
    """Convert a Word paragraph with formatting to ReportLab markup"""
    if not paragraph.runs:
        return escape_html(paragraph.text)
    
    result = ""
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
            
        formatted_text = escape_html(text)
        
        # Apply formatting based on run properties
        if run.bold:
            formatted_text = f"<b>{formatted_text}</b>"
        if run.italic:
            formatted_text = f"<i>{formatted_text}</i>"
        if run.underline:
            formatted_text = f"<u>{formatted_text}</u>"
        
        # Handle font size
        if run.font.size:
            size_pt = int(run.font.size.pt)
            formatted_text = f'<font size="{size_pt}">{formatted_text}</font>'
        
        result += formatted_text
    
    return result

def convert_html_to_pdf_perfect(file_path):
    """Perfect HTML to PDF conversion with all formatting working"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
        from reportlab.lib import colors
        
        pdf_file = file_path.replace('.html', '.pdf')
        
        # Read HTML file
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        print(f"Processing HTML for PDF...")
        
        # Extract and clean body content
        body_content = extract_body_content(html_content)
        print(f"Body content preview: {body_content[:500]}...")
        
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(body_content, 'html.parser')
        
        # Create PDF document
        doc = SimpleDocTemplate(pdf_file, pagesize=letter, 
                              rightMargin=72, leftMargin=72, 
                              topMargin=72, bottomMargin=72)
        
        # Create enhanced styles
        styles = getSampleStyleSheet()
        create_perfect_styles(styles)
        
        story = []
        
        # Process HTML with perfect formatting detection
        process_html_perfect(soup, story, styles)
        
        if not story:
            story.append(Paragraph("No content to display", styles['CustomNormal']))
        
        print(f"Building PDF with {len(story)} elements...")
        doc.build(story)
        print(f"PDF created successfully: {pdf_file}")
        
        return pdf_file
        
    except Exception as e:
        print(f"Detailed PDF error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"PDF conversion failed: {str(e)}")

def convert_html_to_docx_perfect(file_path):
    """Perfect HTML to DOCX conversion with all formatting working"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        docx_file = file_path.replace('.html', '.docx').replace('.txt', '.docx')
        
        # Read and process HTML
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        content = extract_body_content(html_content)
        print(f"Processing HTML for DOCX...")
        
        # Parse with BeautifulSoup
        soup = BeautifulSoup(content, 'html.parser')
        
        # Create document
        doc = Document()
        
        # Process with perfect formatting
        process_html_for_docx_perfect(soup, doc)
        
        if len(doc.paragraphs) == 0:
            doc.add_paragraph("No content to display")
        
        print(f"Saving DOCX: {docx_file}")
        doc.save(docx_file)
        print(f"DOCX created successfully")
        
        return docx_file
        
    except Exception as e:
        print(f"Detailed DOCX error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"DOCX conversion failed: {str(e)}")

def create_perfect_styles(styles):
    """Create perfect styles for ReportLab"""
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.lib import colors
    
    # Base styles with proper font support
    styles.add(ParagraphStyle(
        name='CustomNormal',
        parent=styles['Normal'],
        fontSize=12,
        spaceAfter=6,
        textColor=colors.black,
        leading=16,
        fontName='Helvetica'
    ))
    
    # Size variations
    styles.add(ParagraphStyle(
        name='CustomSmall',
        parent=styles['CustomNormal'],
        fontSize=8
    ))
    
    styles.add(ParagraphStyle(
        name='CustomLarge',
        parent=styles['CustomNormal'],
        fontSize=16
    ))
    
    styles.add(ParagraphStyle(
        name='CustomHuge',
        parent=styles['CustomNormal'],
        fontSize=20
    ))
    
    # Alignment styles
    for align_name, align_val in [('Center', TA_CENTER), ('Right', TA_RIGHT), ('Justify', TA_JUSTIFY)]:
        for size in ['Normal', 'Small', 'Large', 'Huge']:
            styles.add(ParagraphStyle(
                name=f'Custom{size}{align_name}',
                parent=styles[f'Custom{size}'],
                alignment=align_val
            ))
    
    # Headers
    for i in range(1, 7):
        font_size = max(20 - (i * 2), 12)
        styles.add(ParagraphStyle(
            name=f'CustomH{i}',
            parent=styles['CustomNormal'],
            fontSize=font_size,
            fontName='Helvetica-Bold',
            spaceAfter=12,
            spaceBefore=8
        ))

def process_html_perfect(soup, story, styles):
    """Process HTML with perfect formatting detection"""
    from reportlab.platypus import Paragraph, Spacer, PageBreak
    
    # Find all text-containing elements
    all_elements = soup.find_all(True)
    
    for element in all_elements:
        if not element or not element.name:
            continue
            
        # Skip if element is nested inside another we'll process
        if element.parent and element.parent.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote']:
            continue
            
        element_text = element.get_text().strip()
        if not element_text and element.name != 'div':
            continue
        
        try:
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = element.name[1]
                text = create_perfect_reportlab_text(element)
                story.append(Paragraph(text, styles[f'CustomH{level}']))
                
            elif element.name == 'p':
                text = create_perfect_reportlab_text(element)
                if text.strip():
                    style = get_perfect_style(element, styles)
                    story.append(Paragraph(text, style))
                    story.append(Spacer(1, 6))
                    
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li', recursive=False):
                    text = create_perfect_reportlab_text(li)
                    if text.strip():
                        story.append(Paragraph(f"• {text}", styles['CustomNormal']))
                story.append(Spacer(1, 12))
                
            elif element.name == 'blockquote':
                text = create_perfect_reportlab_text(element)
                if text.strip():
                    story.append(Paragraph(f"<i>{text}</i>", styles['CustomNormal']))
                    story.append(Spacer(1, 12))
                    
            elif element.name == 'div':
                if 'page-break' in str(element.get('class', [])):
                    story.append(PageBreak())
                elif element_text:
                    text = create_perfect_reportlab_text(element)
                    if text.strip():
                        style = get_perfect_style(element, styles)
                        story.append(Paragraph(text, style))
                        
        except Exception as e:
            print(f"Error processing element {element.name}: {e}")
            continue

def create_perfect_reportlab_text(element):
    """Create perfectly formatted ReportLab text with all formatting"""
    if not element:
        return ""
    
    def process_node(node, formatting_stack=None):
        if formatting_stack is None:
            formatting_stack = {'bold': False, 'italic': False, 'underline': False, 'strike': False, 'size': None}
        
        result = ""
        
        if hasattr(node, 'contents'):
            for child in node.contents:
                if hasattr(child, 'name') and child.name:
                    # Create new formatting stack for this element
                    new_stack = formatting_stack.copy()
                    
                    # Update formatting based on element
                    classes = str(child.get('class', []))
                    
                    # Detect formatting from tags and classes
                    if child.name in ['strong', 'b'] or 'ql-bold' in classes:
                        new_stack['bold'] = True
                    if child.name in ['em', 'i'] or 'ql-italic' in classes:
                        new_stack['italic'] = True
                    if child.name == 'u' or 'ql-underline' in classes:
                        new_stack['underline'] = True
                    if child.name == 's' or 'ql-strike' in classes:
                        new_stack['strike'] = True
                    
                    # Detect size
                    if 'ql-size-small' in classes:
                        new_stack['size'] = 'small'
                    elif 'ql-size-large' in classes:
                        new_stack['size'] = 'large'
                    elif 'ql-size-huge' in classes:
                        new_stack['size'] = 'huge'
                    
                    # Process child content
                    child_text = process_node(child, new_stack)
                    
                    if child_text.strip():
                        # Apply formatting tags in correct order
                        formatted_text = child_text
                        
                        if new_stack['strike']:
                            formatted_text = f"<strike>{formatted_text}</strike>"
                        if new_stack['underline']:
                            formatted_text = f"<u>{formatted_text}</u>"
                        if new_stack['italic']:
                            formatted_text = f"<i>{formatted_text}</i>"
                        if new_stack['bold']:
                            formatted_text = f"<b>{formatted_text}</b>"
                        
                        # Size formatting
                        if new_stack['size'] == 'small':
                            formatted_text = f'<font size="8">{formatted_text}</font>'
                        elif new_stack['size'] == 'large':
                            formatted_text = f'<font size="16">{formatted_text}</font>'
                        elif new_stack['size'] == 'huge':
                            formatted_text = f'<font size="20">{formatted_text}</font>'
                        
                        result += formatted_text
                    
                elif child.name == 'br':
                    result += "<br/>"
                else:
                    # Text node
                    text = str(child).strip()
                    if text:
                        # Apply current formatting stack
                        formatted_text = escape_html(text)
                        
                        if formatting_stack['strike']:
                            formatted_text = f"<strike>{formatted_text}</strike>"
                        if formatting_stack['underline']:
                            formatted_text = f"<u>{formatted_text}</u>"
                        if formatting_stack['italic']:
                            formatted_text = f"<i>{formatted_text}</i>"
                        if formatting_stack['bold']:
                            formatted_text = f"<b>{formatted_text}</b>"
                        
                        if formatting_stack['size'] == 'small':
                            formatted_text = f'<font size="8">{formatted_text}</font>'
                        elif formatting_stack['size'] == 'large':
                            formatted_text = f'<font size="16">{formatted_text}</font>'
                        elif formatting_stack['size'] == 'huge':
                            formatted_text = f'<font size="20">{formatted_text}</font>'
                        
                        result += formatted_text
        else:
            # Simple text node
            result = escape_html(str(node))
        
        return result
    
    return process_node(element)

def get_perfect_style(element, styles):
    """Get perfect style based on element attributes"""
    classes = str(element.get('class', []))
    
    # Determine size
    size = 'Normal'
    if 'ql-size-small' in classes:
        size = 'Small'
    elif 'ql-size-large' in classes:
        size = 'Large'
    elif 'ql-size-huge' in classes:
        size = 'Huge'
    
    # Determine alignment
    if 'ql-align-center' in classes:
        return styles.get(f'Custom{size}Center', styles[f'Custom{size}'])
    elif 'ql-align-right' in classes:
        return styles.get(f'Custom{size}Right', styles[f'Custom{size}'])
    elif 'ql-align-justify' in classes:
        return styles.get(f'Custom{size}Justify', styles[f'Custom{size}'])
    
    return styles[f'Custom{size}']

def process_html_for_docx_perfect(soup, doc):
    """Process HTML for DOCX with perfect formatting"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Process all elements
    elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'blockquote', 'div'])
    
    for element in elements:
        # Skip nested elements that will be processed by their parents
        if element.parent and element.parent.name in ['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'blockquote']:
            continue
            
        element_text = element.get_text().strip()
        if not element_text and element.name != 'div':
            continue
        
        try:
            if element.name.startswith('h') and element.name[1:].isdigit():
                level = int(element.name[1])
                level = min(level, 9)
                text = element.get_text().strip()
                if text:
                    doc.add_heading(text, level)
                    
            elif element.name == 'p':
                if element_text:
                    paragraph = doc.add_paragraph()
                    apply_perfect_docx_formatting(paragraph, element)
                    apply_docx_alignment_and_size(paragraph, element)
                    
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        paragraph = doc.add_paragraph(style='List Bullet')
                        apply_perfect_docx_formatting(paragraph, li)
                        
            elif element.name == 'blockquote':
                if element_text:
                    paragraph = doc.add_paragraph()
                    apply_perfect_docx_formatting(paragraph, element)
                    for run in paragraph.runs:
                        run.italic = True
                        
            elif element.name == 'div':
                if 'page-break' in str(element.get('class', [])):
                    doc.add_page_break()
                elif element_text:
                    paragraph = doc.add_paragraph()
                    apply_perfect_docx_formatting(paragraph, element)
                    apply_docx_alignment_and_size(paragraph, element)
                    
        except Exception as e:
            print(f"Error processing DOCX element {element.name}: {e}")
            continue

def apply_perfect_docx_formatting(paragraph, element):
    """Apply perfect formatting to DOCX paragraph"""
    from docx.shared import Pt
    
    def process_docx_node(node, paragraph, formatting_stack=None):
        if formatting_stack is None:
            formatting_stack = {'bold': False, 'italic': False, 'underline': False, 'strike': False, 'size': None}
        
        if hasattr(node, 'contents'):
            for child in node.contents:
                if hasattr(child, 'name') and child.name:
                    # Create new formatting stack
                    new_stack = formatting_stack.copy()
                    
                    # Update formatting
                    classes = str(child.get('class', []))
                    
                    if child.name in ['strong', 'b'] or 'ql-bold' in classes:
                        new_stack['bold'] = True
                    if child.name in ['em', 'i'] or 'ql-italic' in classes:
                        new_stack['italic'] = True
                    if child.name == 'u' or 'ql-underline' in classes:
                        new_stack['underline'] = True
                    if child.name == 's' or 'ql-strike' in classes:
                        new_stack['strike'] = True
                    
                    if 'ql-size-small' in classes:
                        new_stack['size'] = 8
                    elif 'ql-size-large' in classes:
                        new_stack['size'] = 16
                    elif 'ql-size-huge' in classes:
                        new_stack['size'] = 20
                    
                    # Process child
                    process_docx_node(child, paragraph, new_stack)
                    
                else:
                    # Text node
                    text = str(child).strip()
                    if text:
                        run = paragraph.add_run(text)
                        
                        # Apply all formatting
                        if formatting_stack['bold']:
                            run.bold = True
                        if formatting_stack['italic']:
                            run.italic = True
                        if formatting_stack['underline']:
                            run.underline = True
                        if formatting_stack['strike']:
                            run.font.strike = True
                        if formatting_stack['size']:
                            run.font.size = Pt(formatting_stack['size'])
        else:
            # Simple text
            text = str(node).strip()
            if text:
                run = paragraph.add_run(text)
                if formatting_stack['bold']:
                    run.bold = True
                if formatting_stack['italic']:
                    run.italic = True
                if formatting_stack['underline']:
                    run.underline = True
                if formatting_stack['strike']:
                    run.font.strike = True
                if formatting_stack['size']:
                    run.font.size = Pt(formatting_stack['size'])
    
    process_docx_node(element, paragraph)

def apply_docx_alignment_and_size(paragraph, element):
    """Apply alignment and size to DOCX paragraph"""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    
    classes = str(element.get('class', []))
    
    # Alignment
    if 'ql-align-center' in classes:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif 'ql-align-right' in classes:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif 'ql-align-justify' in classes:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Size (apply to all runs if not already set)
    size = None
    if 'ql-size-small' in classes:
        size = 8
    elif 'ql-size-large' in classes:
        size = 16
    elif 'ql-size-huge' in classes:
        size = 20
    
    if size:
        for run in paragraph.runs:
            if not run.font.size:  # Only set if not already set
                run.font.size = Pt(size)

def escape_html(text):
    """Escape HTML special characters"""
    if not text:
        return ""
    return (str(text)
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
            .replace('"', '&quot;')
            .replace("'", '&#39;'))

def extract_body_content(html_content):
    """Extract content from HTML body"""
    try:
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL)
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
        
        body_match = re.search(r'<body[^>]*>(.*?)</body>', html_content, re.DOTALL)
        if body_match:
            return body_match.group(1)
        else:
            return html_content
    except Exception as e:
        print(f"Error extracting body content: {e}")
        return html_content

@app.route('/download/<path:filename>')
def download_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({'error': 'File not found'}), 404

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    app.run(debug=True)